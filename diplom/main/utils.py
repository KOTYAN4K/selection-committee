from bs4 import BeautifulSoup
import requests
from django.shortcuts import get_object_or_404
from docx import Document as docs
from django.core.mail import send_mail

from account.models import User
from diplom.settings import EMAIL_HOST_USER
from main.models import School, Parent, Document, Admission, Applicant, ApplicantAdmissionView

urls = [
    'https://edu.tatar.ru/agryz/type/1',
    'https://edu.tatar.ru/aznakaevo/type/1',
    'https://edu.tatar.ru/aksubaevo/type/1',
    'https://edu.tatar.ru/aktanysh/type/1',
    'https://edu.tatar.ru/alekseevo/type/1',
    'https://edu.tatar.ru/alkeevo/type/1',
    'https://edu.tatar.ru/almet/type/1',
    'https://edu.tatar.ru/apastovo/type/1',
    'https://edu.tatar.ru/arsk/type/1',
    'https://edu.tatar.ru/atnya/type/1',
    'https://edu.tatar.ru/bauly/type/1',
    'https://edu.tatar.ru/baltasi/type/1',
    'https://edu.tatar.ru/bugulma/type/1',
    'https://edu.tatar.ru/buinsk/type/1',
    'https://edu.tatar.ru/v_uslon/type/1',
    'https://edu.tatar.ru/v_gora/type/1',
    'https://edu.tatar.ru/n_chelny/type/1',
    'https://edu.tatar.ru/drozhanoye/type/1',
    'https://edu.tatar.ru/elabuga/type/1',
    'https://edu.tatar.ru/zainsk/type/1',
    'https://edu.tatar.ru/z_dol/type/1',
    'https://edu.tatar.ru/kaybitcy/type/1',
    'https://edu.tatar.ru/k_ustye/type/1',
    'https://edu.tatar.ru/kukmor/type/1',
    'https://edu.tatar.ru/laishevo/type/1',
    'https://edu.tatar.ru/l-gorsk/type/1',
    'https://edu.tatar.ru/mamadysh/type/1',
    'https://edu.tatar.ru/mendeleevsk/type/1',
    'https://edu.tatar.ru/menzelinsk/type/1',
    'https://edu.tatar.ru/muslum/type/1',
    'https://edu.tatar.ru/nkamsk/type/1',
    'https://edu.tatar.ru/nsheshma/type/1',
    'https://edu.tatar.ru/nurlat/type/1',
    'https://edu.tatar.ru/pestretcy/type/1',
    'https://edu.tatar.ru/r_sloboda/type/1',
    'https://edu.tatar.ru/saby/type/1',
    'https://edu.tatar.ru/sarmanovo/type/1',
    'https://edu.tatar.ru/spassk/type/1',
    'https://edu.tatar.ru/tetyushi/type/1',
    'https://edu.tatar.ru/tukaj/type/1',
    'https://edu.tatar.ru/tulachi/type/1',
    'https://edu.tatar.ru/cheremshan/type/1',
    'https://edu.tatar.ru/chistopol/type/1',
    'https://edu.tatar.ru/yutaza/type/1',
]


def parse_schools():
    # Очистите существующие данные
    School.objects.all().delete()
    for url in urls:
        parse_info(url)


def parse_info(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    for schools in soup.find_all('ul', class_='edu-list col-md-4'):
        for school in schools.find_all('li'):
            name = school.get_text()
            School.objects.create(name=name).save()


def make_docunment(student):
    ...


def create_account(applicant):
    if not User.objects.filter(email=applicant.email).exists():
        new_user = User()
        new_user.username = (str(applicant.pk).zfill(2) +
                             str(applicant.birth_date.year) +
                             str(applicant.birth_date.month).zfill(2) +
                             str(applicant.birth_date.day).zfill(2))
        password = User.objects.make_random_password()
        new_user.set_password(password)

        new_user.email = applicant.email
        new_user.student = applicant
        new_user.save()
        applicant.change_status_to_answered()
        applicant.save()
        send_mail(
            "Ваша заявка принята.",
            f"""Ваша заявка принята, приступайте к заполнению вашей личной страницы с документами.
Вот ваши данные для авторизации на сайте:
Логин:{new_user.username}
Пароль:{password}

Можно также использовать почту для авторизации.""",
            EMAIL_HOST_USER,
            [new_user.email],
            fail_silently=False,
        )
        parents = Parent.objects.create(student=applicant)
        documents = Document.objects.create(student=applicant)
        admission = Admission.objects.create(applicant=applicant)
        parents.save()
        documents.save()
        admission.save()
        ApplicantAdmissionView.objects.create(applicant=applicant, admission=admission,
                                              document=documents, parent=parents).save()
        return True
    return False


def confirm_student(admission):
    admission.change_status_to_accepted()
    return True


def deny_student(admission):
    admission.change_status_to_accepted()
    return True


def warn_student(admission):
    admission.change_status_to_accepted()
    return True


def send_invite_email(email, subject, message):
    send_mail(
        subject,
        message,
        EMAIL_HOST_USER,
        [email],
        fail_silently=False,
    )


def replace_text_in_runs(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            full_text = ''.join([run.text for run in paragraph.runs])
            new_text = full_text.replace(key, value)

            for i in range(len(paragraph.runs)):
                paragraph.runs[i].text = ''

            paragraph.runs[0].text = new_text


def fill_template(person_id, template_path):
    person = get_object_or_404(Applicant, id=person_id)
    doc = docs(template_path)

    replacements = {
        '{last_name}': f'{person.first_name}',
        '{first_name}': person.last_name,
        '{patronymic}': person.patronymic,
        '{birth_date}': person.birth_date.strftime('%d.%m.%Y'),
        '{passport}': f'Серия  {person.document.get_passport_series()}   № {person.document.get_passport_num()}',
        '{issued_by}': person.document.issued_by,
        '{issue_date}': person.document.issue_date.strftime('%d.%m.%Y'),
        '{phone}': person.phone,
        '{snils}': person.document.SNILS,
        '{addmissions}': person.student.get_departments(),
    }

    for paragraph in doc.paragraphs:
        replace_text_in_runs(paragraph, replacements)


    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)

    return doc